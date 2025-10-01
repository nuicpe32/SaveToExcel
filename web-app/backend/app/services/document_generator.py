from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
from typing import Dict
import os

class DocumentGenerator:
    def __init__(self, template_dir: str = "/app/templates"):
        self.template_dir = template_dir

    def thai_date(self, date: datetime) -> str:
        thai_months = [
            "มกราคม", "กุมภาพันธ์", "มีนาคม", "เมษายน", "พฤษภาคม", "มิถุนายน",
            "กรกฎาคม", "สิงหาคม", "กันยายน", "ตุลาคม", "พฤศจิกายน", "ธันวาคม"
        ]
        day = date.day
        month = thai_months[date.month - 1]
        year = date.year + 543
        return f"{day} {month} {year}"

    def generate_bank_account_doc(self, data: Dict) -> str:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("บันทึกข้อความ")
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(16)
        run.bold = True

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"ส่วนราชการ   กองปราบปรามการกระทำความผิดเกี่ยวกับอาชญากรรมทางเทคโนโลยี")
        p.add_run("\n")
        p.add_run(f"ที่   {data.get('document_number', '')}")
        p.add_run("\n")
        p.add_run(f"วันที่   {self.thai_date(data.get('request_date', datetime.now()))}")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("เรื่อง   ขอข้อมูลบัญชีธนาคาร")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"เรียน   {data.get('bank_name', '')}")

        doc.add_paragraph()

        content = doc.add_paragraph()
        content.add_run(f"          ตามที่ได้รับแจ้งเรื่องราวการกระทำผิด จึงขอความอนุเคราะห์ข้อมูลบัญชีธนาคาร ")
        content.add_run(f"เลขที่บัญชี {data.get('account_number', '')} ชื่อบัญชี {data.get('account_name', '')} ")
        content.add_run(f"สาขา {data.get('branch', '')} เพื่อประกอบการสืบสวนสอบสวน")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("          จึงเรียนมาเพื่อโปรดพิจารณา")

        doc.add_paragraph()
        doc.add_paragraph()

        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature.add_run("ขอแสดงความนับถือ\n\n")
        signature.add_run("(                                      )\n")
        signature.add_run("พ.ต.ท.\n")
        signature.add_run("รอง ผกก.ปอท.")

        filename = f"bank_account_{data.get('document_number', 'doc')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = os.path.join("/tmp", filename)
        doc.save(filepath)

        return filepath

    def generate_suspect_summons_doc(self, data: Dict) -> str:
        doc = Document()

        section = doc.sections[0]
        section.top_margin = Inches(0.3)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        heading = doc.add_paragraph()
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = heading.add_run("บันทึกข้อความ")
        run.font.name = "TH SarabunPSK"
        run.font.size = Pt(16)
        run.bold = True

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("เรื่อง   หมายเรียกผู้ต้องหา")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run(f"เรียน   {data.get('full_name', '')}")

        doc.add_paragraph()

        content = doc.add_paragraph()
        content.add_run(f"          ด้วยท่านถูกกล่าวหาว่ากระทำผิด{data.get('charge', '')} ")
        content.add_run(f"ในคดีหมายเลขดำที่ {data.get('case_number', '')} ")
        content.add_run(f"จึงขอให้ท่านมาพบเพื่อทำการสอบปากคำเป็นผู้ต้องหา ")
        content.add_run(f"ณ {data.get('summons_location', 'กองปราบปรามการกระทำความผิดเกี่ยวกับอาชญากรรมทางเทคโนโลยี')} ")
        content.add_run(f"ในวันที่ {self.thai_date(data.get('summons_date', datetime.now()))} ")
        content.add_run(f"เวลา {data.get('summons_time', '09.00')} น.")

        doc.add_paragraph()

        p = doc.add_paragraph()
        p.add_run("          หากท่านไม่มาตามนัด จะดำเนินการตามกฎหมายต่อไป")

        doc.add_paragraph()
        doc.add_paragraph()

        signature = doc.add_paragraph()
        signature.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        signature.add_run("(                                      )\n")
        signature.add_run("                พ.ต.ต.\n")
        signature.add_run("              รอง ผกก.ปอท.")

        filename = f"suspect_summons_{data.get('document_number', 'doc')}_{datetime.now().strftime('%Y%m%d%H%M%S')}.docx"
        filepath = os.path.join("/tmp", filename)
        doc.save(filepath)

        return filepath